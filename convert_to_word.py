import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_md_to_docx(md_file, docx_file):
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Convert markdown to HTML
    html_content = markdown.markdown(md_content)

    # Create a new Word document
    doc = Document()

    # Set document properties
    doc.core_properties.title = "E-Library Book Recommender System - Project Report"
    doc.core_properties.author = "Project Team"

    # Add title
    title = doc.add_heading('E-Library Book Recommender System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('Project Report', 1)
    doc.add_paragraph()

    # Process content
    lines = md_content.split('\n')
    current_list_level = 0

    for line in lines:
        # Skip empty lines
        if not line.strip():
            continue

        # Handle headers
        if line.startswith('#'):
            level = len(re.match('^#+', line).group())
            text = line.lstrip('#').strip()
            doc.add_heading(text, level)
            continue

        # Handle lists
        if line.strip().startswith(('- ', '* ', '1. ')):
            if line.strip().startswith(('1. ')):
                doc.add_paragraph(line.strip()[3:], style='List Number')
            else:
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            continue

        # Handle code blocks
        if line.strip().startswith('```'):
            continue

        # Regular paragraph
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.style = 'Normal'

    # Save the document
    doc.save(docx_file)

if __name__ == '__main__':
    convert_md_to_docx('Project_Report.md', 'Project_Report.docx') 